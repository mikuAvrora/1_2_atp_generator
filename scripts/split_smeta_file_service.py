"""
    Нужно разделить разные части сметы на разные файлы.
    В смете может быть несколько разных компаний, которые выполняют разные работы.
    Нужно сделать так, чтобы если есть несколько компаний, то они были в разных файлах.
    Если компаний нет, то все в одном файле.
    Формат сметы файла .docx.
"""


import re

from docx import Document


def get_company_name_from_file_name(file_name) -> list or str:
    """ Получает название компании из названия файла """
    bs_name_re = r"[A-Z]{3}_[A-Za-z]+"
    bs_name = re.findall(bs_name_re, file_name)
    try:
        return bs_name
    except Exception as e:
        return Exception(f"Ошибка при получении названия компании из названия файла: {e}")


def get_saved_files(file):
    """ Получает список сохраненных файлов """
    saved_files = [file]
    return saved_files


def save_file(company_name, doc_part, input_file_path, output_folder):
    """ Сохраняет файл """
    file_path = f"{output_folder}/{company_name}.docx"
    doc_part.save(file_path)
    saved_files = get_saved_files(file_path)
    return saved_files


def split_docx_by_paragraph(input_file_path, output_folder):
    document = Document(input_file_path)
    split_indexes = []
    for i, table in enumerate(document.tables):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    text = re.sub(r"\s+", " ", text)
                    text = text.replace('"', '')
                    if "СВОДНЫЙ СМЕТНЫЙ РАСЧЕТ СТОИМОСТИ СТРОИТЕЛЬСТВА" in text and i not in split_indexes:
                        split_indexes.append(i)
                    elif "СМЕТНЫЙ РАСЧЕТ СТОИМОСТИ СТРОИТЕЛЬСТВА" in text and i not in split_indexes:
                        split_indexes.append(i)
    if len(split_indexes) == 0:
        return Exception("Не найдено ни одной таблицы с названием 'СВОДНЫЙ СМЕТНЫЙ РАСЧЕТ СТОИМОСТИ СТРОИТЕЛЬСТВА'")
    try:
        company_names = get_company_name_from_file_name(input_file_path)
    except Exception as e:
        return Exception(e)
    first_file = None
    for idx in range(len(split_indexes) - 1):
        start_index = split_indexes[idx]
        end_index = split_indexes[idx + 1]
        doc_part = Document()
        for element in document.element.body[start_index:end_index + 1]:
            doc_part.element.body.append(element)
        try:
            first_file = save_file(company_names[idx], doc_part, input_file_path, output_folder)
        except IndexError:
            first_file = save_file(company_names[0], doc_part, input_file_path, output_folder)
    last_part_start_index = split_indexes[-1]
    doc_last_part = Document()
    for element in document.element.body[last_part_start_index:]:
        doc_last_part.element.body.append(element)
    last_file = save_file(company_names[-1], doc_last_part, input_file_path, output_folder)
    return [first_file, last_file]
