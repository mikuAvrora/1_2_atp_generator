import json
import os
import traceback
import re
import tkinter as tk
import json

from tkinter import messagebox, filedialog
from bs4 import BeautifulSoup
from num2words import num2words
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT  # type: ignore

from scripts.models import Project
from datetime import datetime

from scripts.split_smeta_file_service import split_docx_by_paragraph, get_company_name_from_file_name



def get_service_from_config():
    with open('config/config.json', 'r', encoding='utf-8') as file:
        config = json.load(file)
        return config.get('with_service', False)


def log_text(text) -> None:
    global was_error
    was_error = True
    with open('media/logs.txt', 'a', encoding="utf-8") as log_file:
        current_datetime = datetime.now()
        formatted_datetime = current_datetime.strftime('%Y-%m-%d %H:%M:%S')
        log_file.write(f'{formatted_datetime}: {text}\n')



def get_work_folder() -> str | None:
    folder_path: str
    with open("config/work_folder.json", "r") as f:
        folder_path = json.load(f)["folder_path"]
    if folder_path != "":
        return folder_path
    else:
        send_message("Укажите путь к рабочей папке")
        return None


def set_work_folder(folder_path):
    with open("config/work_folder.json", "w") as f:
        json.dump({"folder_path": folder_path}, f, ensure_ascii=False)
    send_message("Новое местоположение рабочей папки: \"" + folder_path + "\"")


def send_message(message):
    root = tk.Tk()
    root.withdraw()
    print(message)
    messagebox.showinfo("Формировать отчет", message)


def browse_folder(entry_var: tk.StringVar) -> None:
    folder_selected = filedialog.askdirectory()
    set_work_folder(folder_selected)
    entry_var.set(folder_selected)


def get_text_with(text, latin=True, kirillica=False, chars=("_", "!")):
    pattern = r'\b(?:' + '|'.join(chars) + r')\b'
    text = re.sub(pattern, ' ', text)

    if latin:
        text = re.sub(r'[^a-zA-Z_]+', ' ', text)
    if kirillica:
        text = re.sub(r'[^а-яА-Я_]+', ' ', text)

    return [word.strip() for word in text.split() if word.strip()]


def get_orders(project: Project) -> dict:
    # открыть папку
    folder = get_work_folder()
    if folder == None:
        return {"status": -1}

    # получить все файлы в папке
    files = os.listdir(folder)

    # найти html файлы
    html_file_path = None
    if files:
        for file in files:
            if file.endswith(('.html')):
                html_file_path = folder + "/" + file
    else:
        send_message("В рабочей папке нет файлов")
        return {"status": -1}

    if html_file_path == None:
        send_message("В папке нет html файла")

    # преобразовать html в правильный формат
    html_file = open(html_file_path, 'r', encoding='utf-8')  # type: ignore
    html_content = replace_p_tags_with_br(html_file.read())

    # преобразовать html 
    soup = BeautifulSoup(html_content, 'html.parser')
    body = soup.find('body')
    result = []

    try:
        dogovor_data = body.findChildren(recursive=False)[0].find_all("table")[0].find(
            "get_dogovor_data").text  # type: ignore
        main_tables = body.findChildren(recursive=False)[0].find_all("table")[1:]  # type: ignore
        vedomost_texts = [row for row in body.findChildren(recursive=False)[0].find("td").get_text().split('\n') if
                          "ВЕДОМОСТЬ исполнения работ" in row]  # type: ignore

        multi_BS_NUMBER = []
        multi_BS_NAME = []
        multi_BS_ADDRESS = []
        multi_ORDER_REGION = []
        multi_ORDER_MANAGER = []
        multi_ORDER_MANAGER_POSITION = []
        multi_TOTAL_SUMM = []
        multi_TOTAL_NDS = []
        multi_TOTAL_SUMM_NDS = []
        multi_TOTAL_SUMM_NDS_WORD = []
        multi_ORDER_NUMBER = []
        multi_ORDER_DATE = []
        multi_TABLE = []

        const_ORDER_DOGOVOR_NUMBER = get_ORDER_DOGOVOR_NUMBER(dogovor_data)  # DONE
        const_ORDER_DOGOVOR_DATE = get_ORDER_DOGOVOR_DATE(dogovor_data)  # DONE

        for text in vedomost_texts:
            multi_BS_NUMBER.append(get_BS_NUMBER(text, f'{html_file}'))  # DONE
            multi_BS_NAME.append(get_BS_NAME(text))  # DONE
            multi_BS_ADDRESS.append(get_BS_ADDRESS(text))  # DONE

        try:
            multi_ORDER_REGION = [regions[iii]["reg_name"] for iii in
                                  [i.text.strip() for i in soup.find_all("region_code")]]  # DONE
            multi_ORDER_MANAGER = [regions[iii]["reg_resp_name"] for iii in
                                   [i.text.strip() for i in soup.find_all("region_code")]]  # DONE
            multi_ORDER_MANAGER_POSITION = [regions[iii]["reg_resp_position"] for iii in
                                            [i.text.strip() for i in soup.find_all("region_code")]]  # DONE
        except KeyError:
            for json_data in regions.items():
                for multi_address in multi_BS_ADDRESS:
                    if re.search(json_data[1]["reg_name"], multi_address, re.IGNORECASE):
                        multi_ORDER_REGION.append(json_data[1]["reg_name"])
                        multi_ORDER_MANAGER.append(json_data[1]["reg_resp_name"])
                        multi_ORDER_MANAGER_POSITION.append(json_data[1]["reg_resp_position"])

        for i in range(1, len(main_tables), 2):
            multi_TOTAL_SUMM.append(get_TOTAL_SUMM(main_tables[i]))  # DONE
            multi_TOTAL_NDS.append(get_TOTAL_NDS(main_tables[i]))  # DONE
            multi_TOTAL_SUMM_NDS.append(get_TOTAL_SUMM_NDS(main_tables[i]))  # DONE

        for i in multi_TOTAL_SUMM_NDS:
            multi_TOTAL_SUMM_NDS_WORD.append(get_TOTAL_SUMM_NDS_WORD(i, num2words(
                int(i.strip().replace(" ", "").replace(",", ".").split(".")[0]), lang='ru'),
                                                                     f'{i.strip().replace(" ", "").replace(",", ".").split(".")[1]}'))
            multi_ORDER_NUMBER.append("")
            multi_ORDER_DATE.append("")

        for i in range(0, len(main_tables), 2):
            multi_TABLE.append(get_TABLE(main_tables[i]))  # DONE

        for i, e in enumerate(multi_TABLE):
            type_of_work = get_TYPE_OF_WORK(html_file_path)[0]
            if len(get_TYPE_OF_WORK(html_file_path)) <= len(multi_TABLE):
                try:
                    type_of_work = get_TYPE_OF_WORK(html_file_path)[i]
                except IndexError:
                    type_of_work = get_TYPE_OF_WORK(html_file_path)[0]

            data = {
                "BS_NUMBER": multi_BS_NUMBER[i] if len(multi_BS_NUMBER) > i else "",
                "BS_NAME": multi_BS_NAME[i] if len(multi_BS_NAME) > i else "",
                "BS_ADDRESS": multi_BS_ADDRESS[i] if len(multi_BS_ADDRESS) > i else "",
                "ORDER_REGION": multi_ORDER_REGION[i] if len(multi_ORDER_REGION) > i else "",
                "ORDER_MANAGER": multi_ORDER_MANAGER[i] if len(multi_ORDER_MANAGER) > i else "",
                "ORDER_NUMBER": multi_ORDER_NUMBER[i] if len(multi_ORDER_NUMBER) > i else "",
                "ORDER_DATE": multi_ORDER_DATE[i] if len(multi_ORDER_DATE) > i else "",
                "TOTAL_SUMM": multi_TOTAL_SUMM[i] if len(multi_TOTAL_SUMM) > i else "",
                "TOTAL_NDS": multi_TOTAL_NDS[i] if len(multi_TOTAL_NDS) > i else "",
                "TOTAL_SUMM_NDS": multi_TOTAL_SUMM_NDS[i] if len(multi_TOTAL_SUMM_NDS) > i else "",
                "TOTAL_SUMM_NDS_WORD": multi_TOTAL_SUMM_NDS_WORD[i] if len(multi_TOTAL_SUMM_NDS_WORD) > i else "",
                "ORDER_DOGOVOR_NUMBER": const_ORDER_DOGOVOR_NUMBER if const_ORDER_DOGOVOR_NUMBER else "",
                "ORDER_DOGOVOR_DATE": const_ORDER_DOGOVOR_DATE if const_ORDER_DOGOVOR_DATE else "",
                "TABLE": multi_TABLE[i] if len(multi_TABLE) > i else "",
                "ORDER_MANAGER_POSITION": multi_ORDER_MANAGER_POSITION[i] if len(multi_ORDER_MANAGER_POSITION) > i else "",
                "TYPE_OF_WORK": type_of_work if type_of_work else "",
            }
            data["RRL_PROLET"] = ["РРЛ пролету ", "РРЛ пролета", ""] if "-" in data['BS_NAME'] else [" ", "БС "]
            if "-" in data['BS_ADDRESS']:
                data["BS_ADDRESSES"] = [ f"{data['BS_NAME'].split('-')[index]} - {data['BS_ADDRESS'].split('-')[index]}" for index, value in enumerate(data["BS_ADDRESS"].split("-"))]
            else:
                data["BS_ADDRESSES"]= [data["BS_ADDRESS"]]
            print('\n')
            print(data["BS_ADDRESSES"])
            print('\n')
            result.append(data)
    except:
        print(traceback.format_exc())
        if project.show_errors_window:
            send_message("Произошла ошибка\n" + traceback.format_exc())
    return {"result": result, "status": 0}


def replace_p_tags_with_br(html_content):
    html_content = html_content.replace("style=\"font-size:0.12in;\"", "")
    html_content = html_content.replace("align=\"left\"", "")
    html_content = html_content.replace("align=\"center\"", "")
    html_content = html_content.replace("<br >", "")
    html_content = html_content.replace("<br>", "")
    html_content = html_content.replace("<p", "<br")
    html_content = html_content.replace("<b", "<br")
    html_content = html_content.replace("</p>", "")
    html_content = html_content.replace("brr", "br")
    html_content = html_content.replace("</br>", "")
    html_content = html_content.replace("</b>", "")
    html_content = html_content.replace("</b>", "")
    html_content = html_content.replace("<center>", "")
    html_content = html_content.replace("</center>", "")
    html_content = html_content.replace("<br >", "")
    html_content = html_content.replace("<br>", "")
    html_content = html_content.replace("<br/>", "")
    html_content = html_content.replace("brody", "body")
    html_content = html_content.replace("\n", "")
    html_content = html_content.replace("Итого стоимость работ", "\nИтого стоимость работ")
    html_content = html_content.replace("Всего общая стоимость работ", "\nВсего общая стоимость работ")
    html_content = html_content.replace("НДС 12%: ", "\nНДС 12%: ")
    html_content = html_content.replace("Номер заказа:", "\nНомер заказа:")
    html_content = html_content.replace("<", "\n<")
    html_content = html_content.replace("Регион: [", "\nРегион: [<region_code>")
    html_content = html_content.replace("] Номер Заявки", "</region_code>]\n Номер Заявки")

    html_content_x = html_content.split("\n")
    for i, e in enumerate(html_content_x):
        if "Итого стоимость работ" in e:
            html_content_x[i] = "<itogo_word>" + e + "</itogo_word>"
        if "Всего общая стоимость работ" in e:
            html_content_x[i] = "<itogo_total_word>" + e + "</itogo_total_word>"
        if "НДС 12%: " in e:
            html_content_x[i] = "<NDC_word>" + e + "</NDC_word>"
        if "к рамочному договору" in e:
            html_content_x[i] = "<get_dogovor_data>" + e + "</get_dogovor_data>"

    html_content = "\n".join(html_content_x)
    return html_content


def get_types_of_works():
    with open('config/work_types.json', 'r', encoding="utf-8") as file:
        return json.load(file)


def get_TYPE_OF_WORK(file_name):
    """ Получить все соответствующие типы работ по имени файла """
    types_of_works = get_types_of_works()
    matches = []
    for key in types_of_works:
        if re.search(key, file_name, re.IGNORECASE):
            matches.append((key, types_of_works[key]))

    matches_sorted = [match[1] for match in sorted(matches, key=lambda x: file_name.lower().index(re.search(x[0], file_name, re.IGNORECASE).group(0).lower()))]

    return matches_sorted


def get_FILE_NAME(ATP_OR_AVR, BS_NAME, TYPE_OF_WORK, index=""):
    if index != "":
        index = " (" + f'{index}' + ")"

    variants = {
        "АТП": {
            "демонтажных работ": f"АТП_ДМР_{BS_NAME}_{index}",
            "монтажных работ": f"АТП_МР_{BS_NAME}_{index}",
            "строительных работ": f"{BS_NAME}_АТП_{index}",
            "электромонтажных работ": f"АТП_ЭМР_{BS_NAME}_{index}",
        },
        "АВР": {
            "демонтажных работ": f"АВР_ДМР_{BS_NAME}_",
            "монтажных работ": f"АВР_МР_{BS_NAME}_",
            "строительных работ": f"АВР_СР_{BS_NAME}_",
            "электромонтажных работ": f"АВР_ЭМР_{BS_NAME}_",
        }
    }

    return variants[ATP_OR_AVR][TYPE_OF_WORK]


def get_regions():
    with open('config/regions.json', 'r', encoding="utf-8") as file:
        return json.load(file)


regions = get_regions()


def get_TABLE(table):
    TABLE = []
    if table:
        rows = table.find_all('tr')
        for row in rows:
            cells = row.find_all('td')
            row_list = []
            for cell in cells:
                row_list.append(cell.text.strip())
            try:
                int(row_list[0])
                i = row_list
                if get_service_from_config():
                    TABLE.append({"N": i[0], "P": i[1], "D": i[2], "ST": i[3], "M": i[4], "C": i[5], "T": i[7], "S": i[6]})
                else:
                    TABLE.append({"N": i[0], "P": i[1], "D": i[2], "M": i[3], "C": i[4], "T": i[6], "S": i[5]})
            except:
                # traceback.print_exc()
                pass
                # print(TABLE)
    return TABLE


def get_BS_NUMBER(text, file_name):
    try:
        return [i for i in file_name.split("_") if "БС№" in i][0]
    except:
        try:
            return [i for i in text.split() if "БС№" in i][0]
        except:
            return ""


def get_BS_NAME(text):
    try:
        words = text.split("\"")
        return words[1]
    except:
        try:
            return get_text_with(text, latin=True, kirillica=False, chars=("_", "!"))[0]
        except:
            try:
                for i in text.split(" "):
                    if "БС" in i:
                        return i

                for i in text.split(" "):
                    if "_" in i:
                        return i
                return text.split(" ")[0]
            except:
                return "Не_известеное_название_БС"


def get_BS_ADDRESS(text):
    return ",".join(text.split("ВЕДОМОСТЬ исполнения работ")[1].split(",")[1:])


def get_ORDER_REGION(soup):
    try:
        comment_elements = soup.find_all(string=lambda text: "Vedomost" in text)
        comments_text = [comment.strip() for comment in comment_elements][0]
        extract_metadata(comments_text)['region']  # type: ignore
        return regions[extract_metadata(comments_text)['region']]["reg_name"]  # type: ignore
    except:
        return "reg_name - Не известено"


def get_ORDER_MANAGER(soup):
    try:
        comment_elements = soup.find_all(string=lambda text: "Vedomost" in text)
        comments_text = [comment.strip() for comment in comment_elements][0]
        extract_metadata(comments_text)['region']  # type: ignore
        return regions[extract_metadata(comments_text)['region']]["reg_resp_name"]  # type: ignore
    except:
        return "reg_resp_name - Не известено"


def get_ORDER_MANAGER_POSITION(soup):
    try:
        comment_elements = soup.find_all(string=lambda text: "Vedomost" in text)
        comments_text = [comment.strip() for comment in comment_elements][0]
        extract_metadata(comments_text)['region']  # type: ignore
        return regions[extract_metadata(comments_text)['region']]["reg_resp_position"]  # type: ignore
    except:
        return "reg_resp_position - Не известено"


def get_ORDER_NUMBER(soup):
    ORDER_NUMBER = ""
    return ORDER_NUMBER


def get_ORDER_DATE(soup):
    ORDER_DATE = ""
    return ORDER_DATE


def get_TOTAL_SUMM(table):
    rows = table.find_all('tr')
    for row in rows:
        cells = row.find_all('td')
        if "Итого:" in cells[0].text.strip():
            return cells[1].text.strip()
    TOTAL_SUMM = ""
    return TOTAL_SUMM


def get_TOTAL_NDS(table):
    rows = table.find_all('tr')
    for row in rows:
        cells = row.find_all('td')
        if "НДС 12%:" in cells[0].text.strip():
            return cells[1].text.strip()
    TOTAL_NDS = ""
    return TOTAL_NDS


def get_TOTAL_SUMM_NDS(table):
    rows = table.find_all('tr')
    for row in rows:
        cells = row.find_all('td')
        if "учетом НДС:" in cells[0].text.strip():
            return cells[1].text.strip()
    TOTAL_SUMM_NDS = ""
    return TOTAL_SUMM_NDS


def get_TOTAL_SUMM_NDS_WORD(num, num_word, kopeiki):
    TOTAL_SUMM_NDS_WORD = f"Всего общая стоимость работ: {num} ( {num_word} ) тенге и {kopeiki} тиын"

    return TOTAL_SUMM_NDS_WORD


def get_ORDER_DOGOVOR_NUMBER(dogovor_data):
    xx = dogovor_data.split()
    for i in xx:
        if "№" in i and "№" != i:
            return i[1:]
    ORDER_DOGOVOR_NUMBER = ""
    return ORDER_DOGOVOR_NUMBER


def get_ORDER_DOGOVOR_DATE(dogovor_data):
    try:
        for i in dogovor_data.split():
            if "г" == i[len(i) - 1]:
                return i[:-1]
    except:
        # traceback.print_exc()
        pass
    ORDER_DOGOVOR_DATE = ""
    return ORDER_DOGOVOR_DATE


def get__there_should_be_an_smeta_if_there_is_this_text():
    there_should_be_an_smeta_if_there_is_this_text: str
    with open("config/config.json", "r", encoding="utf-8") as f:
        there_should_be_an_smeta_if_there_is_this_text = json.load(f)["there_should_be_an_smeta_if_there_is_this_text"]
    if there_should_be_an_smeta_if_there_is_this_text != "":
        return there_should_be_an_smeta_if_there_is_this_text
    else:

        # send_message("Укажите путь к рабочей папке")
        return ""


def get_have_smeta(order) -> bool:
    order_str = f"{order}"
    there_should_be_an_smeta_if_there_is_this_text = get__there_should_be_an_smeta_if_there_is_this_text()
    if there_should_be_an_smeta_if_there_is_this_text.lower() in order_str.lower():
        return True
    return False


def combine_docx(file1, file2, output_file, is_second=False, is_atp=False):
    doc1 = Document(file1)
    doc2 = Document(file2)

    docX = doc1

    docX.add_paragraph('')
    docX.add_paragraph('')
    # docX.add_page_break()

    for section in docX.sections:
        section.orientation = WD_ORIENT.PORTRAIT

    for element in doc2.element.body:
        docX.element.body.append(element)

    combined_doc = docX

    for section in combined_doc.sections:
        section.page_width = Pt(510)
        section.page_height = Pt(728)

    # if len(combined_doc.tables) > 3:
    for index, table in enumerate(combined_doc.tables):
        table.autofit = True  # Отключаем автонастройку ширины столбцов

        if index > 3:
            col_lens = []
            for row in table.rows:
                col_lens.append(len(row.cells))

                num_cols = len(row.cells)
                for cell in row.cells:
                    cell.width = Pt(728 / num_cols)

    combined_doc.save(output_file)

    doc1.save(file1[:-5] + ".docx")


def get_smeta(order):
    # открыть папку

    folder = get_work_folder()
    if folder is None:
        return {"status": -1}

    # получить все файлы в папке
    files = os.listdir(folder)

    # найти docx файлы
    not_used_smeta_files = []
    companies = []

    if files:
        for file in files:
            if file.endswith('.docx'):
                if "Смета" in str(file) or "смета" in str(file) or "заказ" in str(file.lower()):
                    not_used_smeta_files.append(f'{folder}/{file}')
                    companies.append(get_company_name_from_file_name(file))

    splited_files = []

    if len(companies) >= 1 and len(companies[0]) > 1:
        for file in not_used_smeta_files:
            if "смета" in file.lower() and file.endswith('.docx'):
                splited_files = split_docx_by_paragraph(file, folder)
                print(splited_files)
                splited_files = [file for file in splited_files if file ]
                print(splited_files)
                break

    if len(not_used_smeta_files) == 0 and get_have_smeta(order) is True:
        send_message("Для заказа требуется смета которую не нашел в папке. Пожалуйста добавьте смету в папку")
        return []

    if len(splited_files) == 0:
        return not_used_smeta_files

    return [file for sublist in splited_files for file in sublist]


def ADD_END(typez, input_path, output_path, data):
    template = None
    if typez == "avr":
        template = DocxTemplate("templates/ШАБЛОН АВР END.docx")
    else:
        template = DocxTemplate("templates/ШАБЛОН АТП END.docx")

    template.render(data)
    template.save("templates/ШАБЛОН WITH END.docx")

    combine_docx(input_path, "templates/ШАБЛОН WITH END.docx", output_path, True, typez == "АТП")


def create_files(folder, data, tmpl_type, have_smeta=False, index=""):
    if " - " in data['BS_NAME']:
        BS_ADDRESSx = data['BS_ADDRESS']
        BS_ADDRESS = data['BS_ADDRESS'].split(" - ")
        BS_NAME = data['BS_NAME'].split(" - ")
        try:
            data['BS_ADDRESS'] = f'{BS_NAME[0]} - {BS_ADDRESS[0]}\n{BS_NAME[1]} - {BS_ADDRESS[1]}'
        except:
            data['BS_ADDRESS'] = BS_ADDRESSx
    print(data)
    smeta_paths = [""]

    if have_smeta:
        smeta_paths = get_smeta(data) if len(get_smeta(data)) > 0 else [""]

    combined_smeta_files = []

    for smeta_path in smeta_paths:
        if "atp" in tmpl_type:
            if get_service_from_config():
                print("ser")
                template_ATP = DocxTemplate("templates/ШАБЛОН АТП(with_service=True).docx")
            else:
                print("noser")
                template_ATP = DocxTemplate("templates/ШАБЛОН АТП.docx")
            template_ATP.render(data)
            file_name__ATP = get_FILE_NAME("АТП", data['BS_NAME'], data['TYPE_OF_WORK'], index=index)
            output_path = folder + "/" + file_name__ATP + ".docx"
            template_ATP.save(output_path)
            print(smeta_path)
            check= False
            for i in data["BS_NAME"]:
                if i.lower() in smeta_path.lower():
                    check = True
            if smeta_path != "" and check:
                combined_smeta_files.append(smeta_path)
                combine_docx(output_path, smeta_path, output_path, is_second=False, is_atp=True)
                ADD_END("ATP", output_path, output_path, data)
                try:
                    os.remove(smeta_path)
                except Exception as e:
                    print(f"Error deleting file {smeta_path}: {e}")
                break
            elif len(combined_smeta_files) == 0 and len(smeta_paths) > 0 and "" not in smeta_paths:
                bs_name = data['BS_NAME'].split(" – ")
                for bs in bs_name:
                    if bs in smeta_path:
                        combined_smeta_files.append(smeta_path)
                        combine_docx(output_path, smeta_path, output_path, is_second=False, is_atp=True)
                        ADD_END("ATP", output_path, output_path, data)
                        break
            ADD_END("ATP", output_path, output_path, data)
